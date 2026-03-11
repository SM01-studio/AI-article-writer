#!/usr/bin/env python3
"""
HTML to DOCX Conversion Script
将wechat.html转换为微信公众号可上传的docx格式

功能:
1. 精确还原HTML样式（字体、颜色、间距、行高）
2. 自动压缩图片以控制文件大小在14.5MB以内
3. 支持微信公众号常见排版元素

使用方式:
    python html_to_docx.py --input "./output/文章标题/wechat.html" --output "./output/文章标题/wechat.docx"

依赖:
    pip install python-docx Pillow beautifulsoup4 lxml
"""

import argparse
import os
import sys
import io
import re
from pathlib import Path
from typing import Optional, Tuple, List, Dict
from copy import deepcopy

# 尝试导入依赖库
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.table import _Cell
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup, NavigableString
except ImportError:
    print("错误: 需要安装 beautifulsoup4 库")
    print("请运行: pip install beautifulsoup4 lxml")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    print("错误: 需要安装 Pillow 库")
    print("请运行: pip install Pillow")
    sys.exit(1)

# 常量配置
MAX_FILE_SIZE_MB = 14.5  # 微信公众号文件大小限制
MAX_IMAGE_WIDTH = 900    # 最大图片宽度（像素）
DEFAULT_IMAGE_QUALITY = 85  # JPEG默认质量


def parse_style(style_str: str) -> dict:
    """解析内联CSS样式字符串"""
    styles = {}
    if not style_str:
        return styles

    for item in style_str.split(';'):
        item = item.strip()
        if ':' in item:
            key, value = item.split(':', 1)
            styles[key.strip().lower()] = value.strip()

    return styles


def parse_px(value: str) -> float:
    """解析像素值，返回磅值"""
    if not value:
        return 0
    value = value.strip().lower()
    if value.endswith('px'):
        return float(value[:-2]) * 0.75  # px to pt
    elif value.endswith('pt'):
        return float(value[:-2])
    elif value.endswith('em'):
        return float(value[:-2]) * 12  # 假设基准字号12pt
    try:
        return float(value) * 0.75
    except:
        return 0


def parse_color(color_str: str) -> Optional[RGBColor]:
    """解析颜色值"""
    if not color_str:
        return None

    color_str = color_str.strip().lower()

    # 处理十六进制颜色
    if color_str.startswith('#'):
        hex_color = color_str[1:]
        if len(hex_color) == 3:
            hex_color = ''.join([c*2 for c in hex_color])
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        except ValueError:
            return None

    # 处理rgb/rgba
    if color_str.startswith('rgb'):
        match = re.search(r'rgba?\((\d+),\s*(\d+),\s*(\d+)', color_str)
        if match:
            r, g, b = int(match.group(1)), int(match.group(2)), int(match.group(3))
            return RGBColor(r, g, b)

    # 常见颜色名称
    color_map = {
        'black': RGBColor(0, 0, 0),
        'white': RGBColor(255, 255, 255),
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'gray': RGBColor(128, 128, 128),
        'grey': RGBColor(128, 128, 128),
    }

    return color_map.get(color_str)


def compress_image(image_path: str, max_width: int = MAX_IMAGE_WIDTH,
                   quality: int = DEFAULT_IMAGE_QUALITY,
                   max_size_kb: int = None) -> Tuple[io.BytesIO, str]:
    """
    压缩图片并返回字节流
    """
    img = Image.open(image_path)

    # 转换为RGB模式（如果需要）
    if img.mode in ('RGBA', 'P'):
        format_type = 'PNG'
    else:
        format_type = 'JPEG'
        if img.mode != 'RGB':
            img = img.convert('RGB')

    # 调整尺寸
    if img.width > max_width:
        ratio = max_width / img.width
        new_height = int(img.height * ratio)
        img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

    # 保存到字节流
    output = io.BytesIO()

    if format_type == 'JPEG':
        if max_size_kb:
            current_quality = quality
            while current_quality > 50:
                output.seek(0)
                output.truncate()
                img.save(output, format='JPEG', quality=current_quality, optimize=True)
                if output.tell() <= max_size_kb * 1024:
                    break
                current_quality -= 5
            else:
                scale = 0.9
                while output.tell() > max_size_kb * 1024 and scale > 0.3:
                    output.seek(0)
                    output.truncate()
                    new_width = int(img.width * scale)
                    new_height = int(img.height * scale)
                    small_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    small_img.save(output, format='JPEG', quality=current_quality, optimize=True)
                    scale -= 0.1
        else:
            img.save(output, format='JPEG', quality=quality, optimize=True)
    else:
        img.save(output, format='PNG', optimize=True)

    output.seek(0)
    return output, format_type.lower()


def set_paragraph_format(paragraph, styles: dict):
    """设置段落格式"""
    pf = paragraph.paragraph_format

    # 行高 - 使用固定行距或1.5倍行距
    line_height = styles.get('line-height', '')
    if line_height:
        try:
            lh = float(line_height)
            if lh >= 1.8:
                # 1.8倍行距 -> 使用1.5倍行距
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif lh >= 1.5:
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            else:
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 对齐方式
    text_align = styles.get('text-align', '')
    if text_align == 'center':
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif text_align == 'justify':
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif text_align == 'right':
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif text_align == 'left':
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 上下边距 - 限制最大值，避免间距过大
    margin_top = styles.get('margin-top', '')
    margin_bottom = styles.get('margin-bottom', '')

    # 解析margin简写
    if 'margin' in styles:
        margin_parts = styles['margin'].split()
        if len(margin_parts) >= 1:
            if not margin_top:
                margin_top = margin_parts[0]
        if len(margin_parts) >= 2:
            if not margin_bottom:
                margin_bottom = margin_parts[1] if len(margin_parts) == 2 else margin_parts[2]

    # 段前间距 - 限制在6pt以内
    if margin_top:
        pt_val = parse_px(margin_top)
        if pt_val > 0:
            pf.space_before = Pt(min(pt_val, 6))

    # 段后间距 - 限制在8pt以内
    if margin_bottom:
        pt_val = parse_px(margin_bottom)
        if pt_val > 0:
            pf.space_after = Pt(min(pt_val, 8))
    else:
        # 默认段后间距
        pf.space_after = Pt(6)

    # 左右边距/缩进
    padding_left = styles.get('padding-left', styles.get('padding', ''))
    if padding_left:
        pt_val = parse_px(padding_left)
        if pt_val > 0:
            pf.left_indent = Pt(pt_val)


def set_run_format(run, styles: dict):
    """设置文本格式"""
    # 字体大小
    font_size = styles.get('font-size', '')
    if font_size:
        pt_val = parse_px(font_size)
        if pt_val > 0:
            run.font.size = Pt(pt_val)

    # 字体颜色
    color = parse_color(styles.get('color', ''))
    if color:
        run.font.color.rgb = color

    # 粗体
    font_weight = styles.get('font-weight', '')
    if font_weight == 'bold' or font_weight == '700':
        run.font.bold = True

    # 斜体
    font_style = styles.get('font-style', '')
    if font_style == 'italic':
        run.font.italic = True


def add_styled_paragraph(doc, text: str, styles: dict, parent_styles: dict = None):
    """添加带样式的段落"""
    p = doc.add_paragraph()

    # 合并样式
    combined_styles = {}
    if parent_styles:
        combined_styles.update(parent_styles)
    combined_styles.update(styles)

    # 设置段落格式
    set_paragraph_format(p, combined_styles)

    # 添加文本
    run = p.add_run(text)
    set_run_format(run, combined_styles)

    return p


def process_inline_content(doc, element, p: Optional = None, parent_styles: dict = None):
    """处理内联内容（文本、strong等）"""
    styles = parse_style(element.get('style', '')) if element.name else {}
    combined_styles = dict(parent_styles) if parent_styles else {}
    combined_styles.update(styles)

    if p is None:
        p = doc.add_paragraph()
        set_paragraph_format(p, combined_styles)

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                run = p.add_run(text)
                set_run_format(run, combined_styles)
        elif child.name == 'strong':
            strong_styles = parse_style(child.get('style', ''))
            strong_combined = dict(combined_styles)
            strong_combined.update(strong_styles)
            strong_combined['font-weight'] = 'bold'

            for sub_child in child.children:
                if isinstance(sub_child, NavigableString):
                    run = p.add_run(str(sub_child))
                    set_run_format(run, strong_combined)
                else:
                    run = p.add_run(sub_child.get_text())
                    set_run_format(run, strong_combined)
        elif child.name == 'em' or child.name == 'i':
            em_styles = parse_style(child.get('style', ''))
            em_combined = dict(combined_styles)
            em_combined.update(em_styles)
            em_combined['font-style'] = 'italic'

            run = p.add_run(child.get_text())
            set_run_format(run, em_combined)
        elif child.name == 'br':
            run = p.add_run('\n')
        elif child.name == 'a':
            # 链接当作普通文本处理
            run = p.add_run(child.get_text())
            set_run_format(run, combined_styles)
        elif child.name:
            # 其他标签，提取文本
            text = child.get_text()
            if text.strip():
                run = p.add_run(text)
                set_run_format(run, combined_styles)

    return p


def add_left_border_paragraph(doc, text: str, styles: dict, border_color: str = '#6366F1', border_width: int = 4):
    """
    添加带左边框的段落（模拟h2的border-left效果）
    使用单列表格实现
    """
    # 解析颜色
    color = parse_color(border_color)
    if not color:
        color = RGBColor(99, 102, 241)

    # 创建单列表格
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False

    # 获取单元格
    cell = table.cell(0, 0)

    # 设置单元格内容
    cell.text = text

    # 设置段落格式
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            run.font.bold = True
            font_size = styles.get('font-size', '18px')
            pt_val = parse_px(font_size)
            if pt_val > 0:
                run.font.size = Pt(pt_val)
            color_val = parse_color(styles.get('color', '#000'))
            if color_val:
                run.font.color.rgb = color_val
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 设置左边框
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    # 左边框
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), str(border_width * 8))  # 1/8 pt
    left_border.set(qn('w:color'), border_color.replace('#', ''))
    tcBorders.append(left_border)

    # 其他边框设为无
    for edge in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)

    tcPr.append(tcBorders)

    # 设置表格宽度
    table.columns[0].width = Cm(15.5)  # 约等于100%宽度

    # 设置单元格内边距
    cell_margin = OxmlElement('w:tcMar')
    for margin_name, value in [('left', '12'), ('top', '0'), ('bottom', '0'), ('right', '0')]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(int(value) * 20))  # twips
        margin.set(qn('w:type'), 'dxa')
        cell_margin.append(margin)
    tcPr.append(cell_margin)

    # 添加上下间距 - 减小间距
    p_before = doc.add_paragraph()
    p_before.paragraph_format.space_before = Pt(12)
    p_before.paragraph_format.space_after = Pt(0)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

    return table


def process_element(doc, element, base_dir: Path, images_data: list, parent_styles: dict = None):
    """递归处理HTML元素"""
    if element.name is None:
        return

    styles = parse_style(element.get('style', ''))
    combined_styles = dict(parent_styles) if parent_styles else {}
    combined_styles.update(styles)

    if element.name == 'h1':
        # 一级标题
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(element.get_text())
        run.font.bold = True
        font_size = styles.get('font-size', '22px')
        pt_val = parse_px(font_size)
        run.font.size = Pt(pt_val if pt_val > 0 else 22)
        color = parse_color(styles.get('color', '#000'))
        if color:
            run.font.color.rgb = color
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    elif element.name == 'h2':
        # 二级标题 - 使用带左边框的样式
        border_left = styles.get('border-left', '')
        border_color = '#6366F1'
        border_width = 4

        if border_left:
            parts = border_left.split()
            if len(parts) >= 3:
                border_width = int(parts[0].replace('px', ''))
                border_color = parts[2] if parts[1] == 'solid' else border_color

        add_left_border_paragraph(doc, element.get_text(), styles, border_color, border_width)

    elif element.name == 'h3':
        # 三级标题
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(element.get_text())
        run.font.bold = True
        font_size = styles.get('font-size', '16px')
        pt_val = parse_px(font_size)
        run.font.size = Pt(pt_val if pt_val > 0 else 16)
        color = parse_color(styles.get('color', '#333'))
        if color:
            run.font.color.rgb = color
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    elif element.name == 'p':
        # 段落
        p = doc.add_paragraph()
        set_paragraph_format(p, combined_styles)
        process_inline_content(doc, element, p, combined_styles)

    elif element.name == 'div':
        # 检查是否是图片容器
        img = element.find('img', recursive=False)
        if img:
            # 图片
            src = img.get('src', '')
            if src:
                img_path = base_dir / src

                if img_path.exists():
                    try:
                        img_stream, fmt = compress_image(str(img_path))
                        images_data.append((img_stream, fmt))

                        # 添加图片段落
                        p = doc.add_paragraph()
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(8)
                        p.paragraph_format.space_after = Pt(8)

                        run = p.add_run()

                        # 获取图片尺寸
                        img_stream.seek(0)
                        pil_img = Image.open(img_stream)
                        img_width, img_height = pil_img.size
                        img_stream.seek(0)

                        # 计算在文档中的宽度
                        max_doc_width = Inches(5.5)  # 适合微信编辑器
                        if img_width > 5.5 * 96:
                            width = max_doc_width
                        else:
                            width = Inches(img_width / 96)

                        run.add_picture(img_stream, width=width)

                    except Exception as e:
                        print(f"警告: 处理图片失败 {img_path}: {e}")
                else:
                    print(f"警告: 图片不存在: {img_path}")
        else:
            # 递归处理div内的内容
            for child in element.children:
                if child.name:
                    process_element(doc, child, base_dir, images_data, combined_styles)

    elif element.name == 'ul':
        # 无序列表
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            li_styles = parse_style(li.get('style', ''))
            li_combined = dict(combined_styles)
            li_combined.update(li_styles)

            set_paragraph_format(p, li_combined)

            # 处理li内的内容
            for child in li.children:
                if isinstance(child, NavigableString):
                    run = p.add_run(str(child))
                    set_run_format(run, li_combined)
                elif child.name == 'strong':
                    strong_styles = parse_style(child.get('style', ''))
                    strong_combined = dict(li_combined)
                    strong_combined.update(strong_styles)
                    strong_combined['font-weight'] = 'bold'

                    run = p.add_run(child.get_text())
                    set_run_format(run, strong_combined)
                else:
                    run = p.add_run(child.get_text())
                    set_run_format(run, li_combined)

    elif element.name == 'ol':
        # 有序列表
        for idx, li in enumerate(element.find_all('li', recursive=False), 1):
            p = doc.add_paragraph(style='List Number')
            li_styles = parse_style(li.get('style', ''))
            li_combined = dict(combined_styles)
            li_combined.update(li_styles)

            set_paragraph_format(p, li_combined)

            for child in li.children:
                if isinstance(child, NavigableString):
                    run = p.add_run(str(child))
                    set_run_format(run, li_combined)
                elif child.name == 'strong':
                    strong_styles = parse_style(child.get('style', ''))
                    strong_combined = dict(li_combined)
                    strong_combined.update(strong_styles)
                    strong_combined['font-weight'] = 'bold'

                    run = p.add_run(child.get_text())
                    set_run_format(run, strong_combined)
                else:
                    run = p.add_run(child.get_text())
                    set_run_format(run, li_combined)

    elif element.name == 'blockquote':
        # 引用块
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(8)

        for child in element.descendants:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    run = p.add_run(text + ' ')
                    run.font.size = Pt(11)
                    run.font.italic = True
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    elif element.name == 'hr':
        # 分割线
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(16)

        # 添加一条细线
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), 'EEEEEE')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    elif element.name == 'table':
        # 表格
        rows = element.find_all('tr', recursive=False)
        if rows:
            first_row_cells = rows[0].find_all(['th', 'td'], recursive=False)
            num_cols = len(first_row_cells)

            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'

            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'], recursive=False)
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        table_cell = table.rows[i].cells[j]
                        table_cell.text = cell.get_text().strip()

                        if cell.name == 'th' or i == 0:
                            for paragraph in table_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

            doc.add_paragraph()


def html_to_docx(html_path: str, output_path: str, max_size_mb: float = MAX_FILE_SIZE_MB) -> bool:
    """
    将HTML文件转换为DOCX格式
    """
    html_path = Path(html_path)
    output_path = Path(output_path)

    if not html_path.exists():
        print(f"错误: HTML文件不存在: {html_path}")
        return False

    # 读取HTML文件
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    # 解析HTML
    soup = BeautifulSoup(html_content, 'lxml')

    # 创建DOCX文档
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 获取HTML基目录
    base_dir = html_path.parent

    # 存储图片数据
    images_data = []

    # 查找主要内容容器
    main_container = soup.find('section')
    if not main_container:
        main_container = soup.find('article')  # 支持 article 标签
    if not main_container:
        body = soup.find('body')
        if body:
            main_container = body.find('div', style=lambda x: x and 'max-width' in x if x else False)
            if not main_container:
                main_container = body
        else:
            main_container = soup

    if not main_container:
        print("错误: 无法找到有效的内容容器")
        return False

    # 遍历所有直接子元素
    for element in main_container.children:
        if element.name is None:
            continue
        process_element(doc, element, base_dir, images_data)

    # 保存文档
    try:
        temp_output = io.BytesIO()
        doc.save(temp_output)
        file_size_mb = temp_output.tell() / (1024 * 1024)

        if file_size_mb > max_size_mb:
            print(f"警告: 文件大小 ({file_size_mb:.2f}MB) 超过限制 ({max_size_mb}MB)")
            print("建议: 减少图片数量或进一步压缩图片")
        else:
            print(f"文件大小: {file_size_mb:.2f}MB (限制: {max_size_mb}MB)")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'wb') as f:
            f.write(temp_output.getvalue())

        print(f"✅ DOCX文件已保存: {output_path}")
        return True

    except Exception as e:
        print(f"错误: 保存文档失败: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(
        description="将wechat.html转换为微信公众号可上传的docx格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python html_to_docx.py --input "./output/文章标题/wechat.html"
    python html_to_docx.py --input "./output/文章标题/wechat.html" --output "./output/文章标题/wechat.docx"
    python html_to_docx.py --input "./output/文章标题/wechat.html" --max-size 10
        """
    )

    parser.add_argument(
        "--input", "-i",
        required=True,
        help="输入的wechat.html文件路径"
    )

    parser.add_argument(
        "--output", "-o",
        help="输出的docx文件路径（默认：与输入文件同目录，后缀改为.docx）"
    )

    parser.add_argument(
        "--max-size", "-m",
        type=float,
        default=MAX_FILE_SIZE_MB,
        help=f"最大文件大小限制(MB)，默认{MAX_FILE_SIZE_MB}"
    )

    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="显示详细信息"
    )

    args = parser.parse_args()

    input_path = Path(args.input)
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / (input_path.stem + '.docx')

    if args.verbose:
        print(f"输入文件: {input_path}")
        print(f"输出文件: {output_path}")
        print(f"大小限制: {args.max_size}MB")
        print("-" * 50)

    success = html_to_docx(str(input_path), str(output_path), args.max_size)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
